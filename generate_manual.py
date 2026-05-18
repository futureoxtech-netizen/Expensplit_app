"""
Build a polished Word manual for end users of "Expense Divider".
No technical terms. Friendly, step-by-step.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Expense Divider - User Manual.docx"

# Brand colours
PRIMARY = RGBColor(0x6C, 0x5C, 0xE7)
ACCENT = RGBColor(0x00, 0xB8, 0x94)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
MUTED = RGBColor(0x6B, 0x6B, 0x7A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def title_block(doc: Document):
    # Coloured "hero" using a 1-row, 1-col table
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Cm(16)
    _shade(cell, "6C5CE7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Expense Divider")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = WHITE
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(28)
    run2 = p2.add_run("User Manual  ·  How to track and split expenses with friends")
    run2.font.size = Pt(13)
    run2.font.color.rgb = WHITE

    doc.add_paragraph()  # spacer


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = PRIMARY
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = DARK
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = MUTED


def body(doc: Document, text: str, *, bold_first_word: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_first_word:
        first, _, rest = text.partition(" ")
        r = p.add_run(first + " ")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = DARK
        r2 = p.add_run(rest)
        r2.font.size = Pt(11)
        r2.font.color.rgb = DARK
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.color.rgb = DARK


def bullet(doc: Document, text: str, *, lead: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if lead:
        r = p.add_run(lead)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = DARK
        r2 = p.add_run(" — " + text)
        r2.font.size = Pt(11)
        r2.font.color.rgb = DARK
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.color.rgb = DARK


def step(doc: Document, n: int, title: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    badge = p.add_run(f"Step {n}  ")
    badge.bold = True
    badge.font.size = Pt(11)
    badge.font.color.rgb = ACCENT
    t = p.add_run(title)
    t.bold = True
    t.font.size = Pt(12)
    t.font.color.rgb = DARK
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.9)
    p2.paragraph_format.space_after = Pt(6)
    r = p2.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = DARK


def callout(doc: Document, title: str, text: str, fill_hex: str = "F1EFFF"):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Cm(16)
    _shade(cell, fill_hex)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = PRIMARY
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK
    doc.add_paragraph()


def build():
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.8)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(11)

    title_block(doc)

    # ──────── Welcome ────────
    heading(doc, "Welcome to Expense Divider")
    body(
        doc,
        "Expense Divider is the easiest way to keep track of who paid for what when "
        "you spend money with other people — trips, roommates, family, the office, "
        "a dinner out, anything. Add what you spent, the app does the math, and "
        "everyone always knows where they stand.",
    )
    body(
        doc,
        "You can also use it just for yourself, to see where your money is going "
        "each week or month, and to get a nicely formatted report you can save or share.",
    )

    callout(
        doc,
        "What is the app for?",
        "1) Splitting bills with friends, family, roommates and trip mates — fairly and "
        "without arguments.\n2) Tracking your personal spending and getting beautiful "
        "monthly reports.",
    )

    # ──────── Getting started ────────
    heading(doc, "Getting started")
    step(doc, 1, "Create your account",
         "Open the app and choose “Create account”. Enter your name, email, a "
         "password, and pick the currency you usually spend in (you can change "
         "this later).")
    step(doc, 2, "Sign in next time",
         "Use your email and password. The app remembers you on your device "
         "so you don't have to type them in every time.")
    step(doc, 3, "Take a look around",
         "There are four sections at the bottom of the screen: Home, Groups, "
         "Activity and Profile. Tap each one to explore.")

    # ──────── Groups ────────
    heading(doc, "Groups — where the magic happens")
    body(
        doc,
        "A group is a shared space where you and other people can add expenses "
        "together. It could be a weekend trip, your apartment, an office team, "
        "a family, or one specific event.",
    )

    heading(doc, "Create a group", level=2)
    step(doc, 1, "Tap “New group”",
         "Go to the Groups tab and tap “New group” at the top.")
    step(doc, 2, "Give it a name",
         "For example: Weekend Trip, Roommates, Office Lunch, Birthday Party.")
    step(doc, 3, "Pick a category and a cover colour",
         "Just for personality — it'll show up in lists and on cards.")
    step(doc, 4, "Choose the group's currency",
         "Expenses in this group will be tracked in that currency.")
    step(doc, 5, "Save",
         "You're now the group owner. You can edit any of these later.")

    heading(doc, "Add people to the group", level=2)
    body(doc, "Two easy ways:")
    bullet(doc, "Open the group, tap the small QR icon at the top right, and either share the code or let them scan the QR with their phone.", lead="By invite code or QR")
    bullet(doc, "Open the group, go to the Members tab, tap “Invite member”, and type their email (they need to be signed up to Expense Divider).", lead="By email")

    # ──────── Adding expenses ────────
    heading(doc, "Adding an expense")
    body(
        doc,
        "Whenever someone in the group pays for something — groceries, fuel, "
        "rent, a meal, anything — record it so the app can keep everything fair.",
    )
    step(doc, 1, "Open the group and tap “Add expense”",
         "The big rounded button at the bottom of the group.")
    step(doc, 2, "What was it and how much",
         "Type a short description (e.g., “Groceries”) and the amount.")
    step(doc, 3, "Pick a category",
         "Food, Groceries, Transport, Rent, Travel and more — pick whichever fits.")
    step(doc, 4, "Who paid?",
         "Tap the person who actually paid for it. By default it's you, but it "
         "could be any group member.")
    step(doc, 5, "Choose a split mode",
         "How should this expense be divided? See the next section.")
    step(doc, 6, "Save",
         "Done. Everyone in the group sees it instantly.")

    heading(doc, "The four ways to split", level=2)
    bullet(doc, "Divide the total evenly across everyone you tick. The most common choice.", lead="Equal")
    bullet(doc, "Type the exact amount each person should pay. The app checks that the parts add up to the total.", lead="Exact amounts")
    bullet(doc, "Type a percentage for each person. The app checks the percentages add up to 100.", lead="Percentage")
    bullet(doc, "Give each person a “share weight”, like 1, 1, 2 — useful when someone should pay double, for example.", lead="Shares")

    callout(
        doc,
        "Don't worry about doing math",
        "Whichever mode you pick, the app shows you a live summary while you "
        "type, so you can see if the parts add up before saving.",
    )

    # ──────── Balances + settle ────────
    heading(doc, "Who owes whom — and settling up")
    body(
        doc,
        "Tap the Balances tab inside any group. You'll see two things:",
    )
    bullet(doc, "How much each person is currently owed or owes overall in the group.", lead="Net balances")
    bullet(doc, "The smallest possible set of payments that would clear everyone's debts. Instead of six awkward transfers, often only one or two are needed.", lead="Suggested payments")
    body(
        doc,
        "When someone actually pays you back (in cash, bank transfer, anything), "
        "open Balances and tap “Settle” on the matching row. The app records "
        "the payment and the balances update for everyone in real time.",
    )

    # ──────── Reports ────────
    heading(doc, "Reports — see where your money goes")
    body(
        doc,
        "From the Home screen tap the big purple “Total tracked” card. You're "
        "taken to Reports, where you can:",
    )
    bullet(doc, "Pick any time range — Today, This week, This month, This year, or a custom range.")
    bullet(doc, "See the total amount you spent and how many transactions that includes.")
    bullet(doc, "View a colourful donut chart of where your money went, by category. Tap a slice to focus on it.")
    bullet(doc, "Read a quick “insight” line, like your top category and average spend per day.")
    bullet(doc, "Tap the big orange button to download a clean, branded PDF report of the chosen period — perfect for keeping records, expense claims, or sharing.")

    # ──────── Notifications & activity ────────
    heading(doc, "Activity & notifications")
    body(
        doc,
        "The bell icon on the Home screen and the Activity tab at the bottom "
        "show everything that's happening in your groups — who added what, who "
        "paid whom. A red number badge tells you how many new things you "
        "haven't seen. Open the Activity tab and the badge resets to zero.",
    )

    # ──────── Profile ────────
    heading(doc, "Your profile")
    body(
        doc,
        "Open the Profile tab to:",
    )
    bullet(doc, "See your name, email and your unique referral code.")
    bullet(doc, "Switch between Light, Dark, and System theme — the app changes instantly.")
    bullet(doc, "Change your default currency — useful if you move country or want to track in a different currency.")
    bullet(doc, "Sign out.")

    # ──────── Tips ────────
    heading(doc, "Tips for getting the most out of it")
    bullet(doc, "Add expenses as they happen — even a small ₨ 100 chai. It only takes a few seconds and saves the “wait, who paid?” conversation later.")
    bullet(doc, "Use categories. They power the colourful donut chart in your monthly Reports.")
    bullet(doc, "Use the “Shares” split when the fair split isn't equal — like when two people share a room and one person gets a private one.")
    bullet(doc, "Send the PDF report to yourself or your accountant at the end of every month.")
    bullet(doc, "Invite people by QR code in person — it's the fastest way at the start of a trip.")

    # ──────── FAQ ────────
    heading(doc, "Frequently asked questions")

    heading(doc, "Do all my friends need an account?", level=2)
    body(
        doc,
        "Anyone you want to be a real member of a group does, yes. Creating an "
        "account is free and takes under a minute.",
    )

    heading(doc, "What if someone doesn't pay me back?", level=2)
    body(
        doc,
        "Expense Divider keeps a clear, honest record of who owes whom — but it "
        "doesn't actually move money. You can show them the Balances screen, "
        "and once they pay you in cash or by transfer, tap “Settle” to mark "
        "the payment.",
    )

    heading(doc, "Can I use different currencies?", level=2)
    body(
        doc,
        "Yes. Each group has its own currency. You also have your personal "
        "default currency in Profile.",
    )

    heading(doc, "Is my data private?", level=2)
    body(
        doc,
        "Yes. Only the people you invite to a group can see that group's "
        "expenses. Your account details, like email and password, are never "
        "shared with anyone.",
    )

    heading(doc, "Does it work on a computer too?", level=2)
    body(
        doc,
        "Yes — it works on your phone and in a web browser. The screens look "
        "and feel the same.",
    )

    # ──────── Closing ────────
    heading(doc, "That's it")
    body(
        doc,
        "You're all set. Create your first group, add an expense, and let "
        "Expense Divider handle the rest. Enjoy keeping money fair, simple, "
        "and stress-free.",
    )

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
